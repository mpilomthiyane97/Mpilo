import os
import pandas as pd
from docx import Document as DocxDocument
from llama_index.core.schema import Document
from llama_index.readers.file import PyMuPDFReader

def load_documents_from_folder(folder_path):
    all_documents = []
    pdf_loader = PyMuPDFReader()

    for filename in os.listdir(folder_path):
        filepath = os.path.join(folder_path, filename)
        ext = os.path.splitext(filename)[1].lower()

        if ext == ".pdf":
            print(f"Loading PDF: {filename}")
            docs = pdf_loader.load(file_path=filepath)
            for d in docs:
                d.metadata['file_name'] = filename
            all_documents.extend(docs)

        elif ext == ".txt":
            print(f"Loading TXT: {filename}")
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
                all_documents.append(Document(text=text, metadata={"file_name": filename}))

        elif ext == ".docx":
            print(f"Loading DOCX: {filename}")
            doc = DocxDocument(filepath)
            text = "\n".join([para.text for para in doc.paragraphs])
            all_documents.append(Document(text=text, metadata={"file_name": filename}))

        elif ext == ".csv":
            print(f"Loading CSV: {filename}")
            df = pd.read_csv(filepath)
            text = "\n".join(df.astype(str).apply(" | ".join, axis=1))
            all_documents.append(Document(text=text, metadata={"file_name": filename}))

        else:
            print(f"Unsupported file type: {filename}")

    return all_documents
